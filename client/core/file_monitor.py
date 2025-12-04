import os
import time
import queue
from PyQt6.QtCore import QThread, pyqtSignal


class FileMonitor(QThread):
    """
    多源监控线程 (Global Cookie Sync Version + Remove Support)
    """
    stats_updated = pyqtSignal(int, int, int)

    def __init__(self):
        super().__init__()
        self.running = True
        self.task_queue = queue.Queue()

        # 监控源列表
        self.sources = []
        self.max_sources = 10

        # 全局 Cookie 缓存
        self.shared_cookies = []

        self.start_time = time.time()
        self.total_initial_sum = 0

        self.last_autosave_time = time.time()
        self.autosave_interval = 60

    def add_source(self, path_or_url, is_web=False):
        self.task_queue.put({
            'type': 'add',
            'path': path_or_url,
            'is_web': is_web
        })
        return True

    def remove_source(self, path):
        """新增：移除监控源"""
        self.task_queue.put({
            'type': 'remove',
            'path': path
        })

    # =========================================================
    #  后台线程方法 (Thread Safe Zone)
    # =========================================================

    def _create_new_driver(self, index):
        try:
            from selenium import webdriver
            from selenium.webdriver.edge.options import Options as EdgeOptions
            from selenium.webdriver.edge.service import Service as EdgeService
        except ImportError:
            return None

        try:
            options = EdgeOptions()
            options.add_argument("--disable-gpu")
            options.add_argument("--no-sandbox")
            options.add_experimental_option('excludeSwitches', ['enable-logging'])
            options.add_experimental_option("detach", True)

            base_path = os.path.join(os.environ['LOCALAPPDATA'], 'InkSprint', 'EdgeData')
            profile_path = os.path.join(base_path, f"Profile_{index}")
            if not os.path.exists(profile_path):
                os.makedirs(profile_path, exist_ok=True)

            options.add_argument(f"user-data-dir={profile_path}")

            log_path = "NUL" if os.name == 'nt' else "/dev/null"
            service = EdgeService(log_output=log_path)

            driver = webdriver.Edge(options=options, service=service)
            return driver
        except Exception as e:
            print(f"[Driver Create Error] {e}")
            return None

    def _inject_cookies_if_available(self, driver):
        if not self.shared_cookies: return False
        try:
            driver.get("https://docs.qq.com/desktop")
            for cookie in self.shared_cookies:
                cookie_dict = {
                    'name': cookie.get('name'),
                    'value': cookie.get('value'),
                    'domain': cookie.get('domain'),
                    'path': cookie.get('path', '/'),
                    'secure': cookie.get('secure', False)
                }
                if 'qq.com' in str(cookie_dict.get('domain', '')):
                    try:
                        driver.add_cookie(cookie_dict)
                    except:
                        pass
            print("[Monitor] 已尝试注入全局 Cookie")
            return True
        except Exception as e:
            print(f"[Cookie Inject Error] {e}")
            return False

    def _update_global_cookies(self, driver):
        try:
            cookies = driver.get_cookies()
            if len(cookies) > 0:
                self.shared_cookies = cookies
        except:
            pass

    def _handle_add_source(self, path, is_web):
        for src in self.sources:
            if src['path'] == path: return

        new_index = len(self.sources)
        new_source = {
            'path': path,
            'type': 'web' if is_web else 'local',
            'initial': 0, 'current': 0, 'is_calibrated': False,
            'mtime': 0, 'driver': None
        }

        if is_web:
            driver = self._create_new_driver(new_index)
            if driver:
                try:
                    self._inject_cookies_if_available(driver)
                    driver.get(path)
                    new_source['driver'] = driver
                    print(f"[Monitor] Web源已启动: {path}")
                except:
                    try:
                        driver.quit()
                    except:
                        pass
                    return
            else:
                return

        self.sources.append(new_source)

    def _handle_remove_source(self, path):
        """处理移除逻辑"""
        for i, src in enumerate(self.sources):
            if src['path'] == path:
                # 1. 如果是 Web，关闭浏览器
                if src['type'] == 'web' and src['driver']:
                    try:
                        src['driver'].quit()
                    except:
                        pass

                # 2. 修正总初始值 (防止移除后增量突变)
                # 逻辑：移除源后，它的 initial 不再参与 total_initial_sum
                # 同时它的 current 也不再参与 total_current_sum
                if src['is_calibrated']:
                    self.total_initial_sum -= src['initial']

                self.sources.pop(i)
                print(f"[Monitor] 已移除源: {path}")
                return

    def _get_web_count(self, driver):
        if not driver: return -1
        try:
            try:
                _ = driver.current_url
            except:
                return -1

            js_script = """
                let text = document.body.innerText;
                let m1 = text.match(/(\\d+)\\s*个?字/);
                if (m1) return parseInt(m1[1]);
                let statusBar = document.querySelector('.word-count-info, .statusbar-simple-text');
                if (statusBar) {
                    let m2 = statusBar.innerText.match(/(\\d+)/);
                    if (m2) return parseInt(m2[1]);
                }
                let m3 = text.match(/Word Count[:：]\\s*(\\d+)/);
                if (m3) return parseInt(m3[1]);
                return -1;
            """
            val = driver.execute_script(js_script)
            if isinstance(val, int) and val >= 0:
                return val
            return -1
        except:
            return -1

    def _get_local_count(self, path):
        if not os.path.exists(path): return 0
        try:
            import docx
        except ImportError:
            return 0
        try:
            ext = os.path.splitext(path)[1].lower()
            content = ""
            if ext == '.docx':
                doc = docx.Document(path)
                parts = [p.text for p in doc.paragraphs]
                for t in doc.tables:
                    for r in t.rows:
                        for c in r.cells:
                            parts.append(c.text)
                content = "".join(parts)
            elif ext == '.txt':
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            return len(content.replace('\n', '').replace(' ', '').replace('\t', '').replace('\r', ''))
        except:
            return 0

    def _trigger_autosave(self):
        try:
            import win32com.client as win32
            import pythoncom
            pythoncom.CoInitialize()
            try:
                word_app = win32.GetActiveObject("Word.Application")
                for doc in word_app.Documents:
                    if not doc.Saved: doc.Save()
            except:
                pass
            finally:
                pythoncom.CoUninitialize()
        except:
            pass

    def run(self):
        print("[Monitor] 线程启动 (Remove Support)")

        while self.running:
            while not self.task_queue.empty():
                try:
                    task = self.task_queue.get_nowait()
                    if task['type'] == 'add':
                        self._handle_add_source(task['path'], task['is_web'])
                    elif task['type'] == 'remove':
                        self._handle_remove_source(task['path'])
                except queue.Empty:
                    break

            total_current_sum = 0

            for src in self.sources:
                if src['type'] == 'web':
                    if src['driver']:
                        val = self._get_web_count(src['driver'])
                        if val != -1:
                            self._update_global_cookies(src['driver'])
                            if not src['is_calibrated']:
                                if val > 0:
                                    src['initial'] = val
                                    src['current'] = val
                                    src['is_calibrated'] = True
                                    self.total_initial_sum += val
                                    print(f"[Monitor] Web校准完成: {val}")
                            else:
                                src['current'] = val

                elif src['type'] == 'local':
                    if os.path.exists(src['path']):
                        mtime = os.path.getmtime(src['path'])
                        if mtime != src['mtime']:
                            val = self._get_local_count(src['path'])
                            if not src['is_calibrated']:
                                src['initial'] = val
                                src['current'] = val
                                src['is_calibrated'] = True
                                self.total_initial_sum += val
                            else:
                                src['current'] = val
                            src['mtime'] = mtime

                if src['is_calibrated']:
                    total_current_sum += src['current']

            total_increment = total_current_sum - self.total_initial_sum

            elapsed = time.time() - self.start_time
            wph = int((total_increment / elapsed) * 3600) if elapsed > 1 else 0

            self.stats_updated.emit(total_current_sum, total_increment, wph)

            if time.time() - self.last_autosave_time > self.autosave_interval:
                self._trigger_autosave()
                self.last_autosave_time = time.time()

            time.sleep(1)

    def stop(self):
        self.running = False
        for src in self.sources:
            if src.get('driver'):
                try:
                    src['driver'].quit()
                except:
                    pass
        self.wait()